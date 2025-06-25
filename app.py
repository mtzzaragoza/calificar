import streamlit as st
import pandas as pd
from docx import Document
import io
import re
import random
import streamlit.components.v1 as components

# Intentar importar librerías de PDF
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configuración de la página
st.set_page_config(page_title="Sistema de Retroalimentación", layout="wide")

# Sidebar para navegación
st.sidebar.title("📚 Sistema de Retroalimentación")
menu_option = st.sidebar.selectbox(
    "Selecciona una opción:",
    ["R3MD - Conjuntos", "R4MD - Proposiciones Lógicas", "R7MD - Mensajes Predefinidos"]
)

# ==================== FUNCIONES COMPARTIDAS ====================

def copy_to_clipboard_js(text):
    """Genera JavaScript para copiar texto al portapapeles"""
    js_code = f"""
    <script>
        navigator.clipboard.writeText(`{text.replace('`', '\\`')}`).then(function() {{
            console.log('Texto copiado al portapapeles');
        }});
    </script>
    """
    return js_code

# ==================== R3MD - CONJUNTOS ====================

def extraer_texto_pdf(pdf_file):
    """Extrae texto de un archivo PDF usando pdfplumber"""
    if not PDF_AVAILABLE:
        raise Exception("Las librerías de PDF no están instaladas. Instala: pip install PyPDF2 pdfplumber")
    
    try:
        texto_completo = ""
        with pdfplumber.open(pdf_file) as pdf:
            for pagina in pdf.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
        return texto_completo.strip()
    except Exception as e:
        try:
            pdf_file.seek(0)
            texto_completo = ""
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for pagina in pdf_reader.pages:
                texto_completo += pagina.extract_text() + "\n"
            return texto_completo.strip()
        except Exception as e2:
            raise Exception(f"Error con pdfplumber: {str(e)} | Error con PyPDF2: {str(e2)}")

def extraer_texto_docx(docx_file):
    """Extrae texto de un archivo DOCX"""
    doc = Document(docx_file)
    texto_parrafos = [p.text for p in doc.paragraphs]
    texto_tablas = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto_tablas.append(cell.text)
    return " ".join(texto_parrafos + texto_tablas).strip()

def extraer_conjunto(texto):
    match_llaves = re.search(r"{([^}]*)}", texto)
    match_corchetes = re.search(r"\[([^\]]*)\]", texto)
    
    contenido = None
    if match_llaves:
        contenido = match_llaves.group(1).strip()
    elif match_corchetes:
        contenido = match_corchetes.group(1).strip()
    else:
        numeros_match = re.search(r"(\d+(?:\s*,\s*\d+)+)", texto)
        if numeros_match:
            contenido = numeros_match.group(1).strip()
        else:
            numero_solo = re.search(r"\b(\d+)\b", texto)
            if numero_solo:
                contenido = numero_solo.group(1).strip()
    
    if contenido:
        numeros = re.findall(r"\d+", contenido)
        numeros_normalizados = [str(int(num)) for num in numeros]
        return set(numeros_normalizados)
    
    return set()

def extraer_nombre(texto):
    match = re.search(r"(?i)nombre completo:\s*(\w+)", texto)
    return match.group(1) if match else "Alumno"

def normalizar_expresion(expresion):
    expr_normalizada = expresion.replace(" ", "")
    normalizaciones = {
        'Ս': '∪',
        'Ո': '∩',
        '´': '′',
        '--': '–',
        '-': '–',
        "'": '′'
    }
    for original, reemplazo in normalizaciones.items():
        expr_normalizada = expr_normalizada.replace(original, reemplazo)
    return expr_normalizada.lower()

def extraer_expresion_y_conjunto(expresion_completa):
    if '=' in expresion_completa:
        partes = expresion_completa.split('=', 1)
        expresion = normalizar_expresion(partes[0].strip())
        conjunto = extraer_conjunto(partes[1].strip())
        return expresion, conjunto
    return "", set()

def buscar_expresion_completa(texto_completo, indice_inciso, expresion_esperada):
    letras = "abcdefghijklmnopqrstuvwxyz"
    letra_inciso = letras[indice_inciso]
    
    expresion_esperada_norm, conjunto_esperado = extraer_expresion_y_conjunto(expresion_esperada)
    
    if not expresion_esperada_norm or not conjunto_esperado:
        return False, ""
    
    encontrado, linea = buscar_por_inciso_exacto(texto_completo, letra_inciso, expresion_esperada_norm, conjunto_esperado)
    if encontrado:
        return True, linea
    
    encontrado, linea = buscar_por_expresion_flexible(texto_completo, expresion_esperada_norm, conjunto_esperado)
    if encontrado:
        return True, linea
    
    return False, ""

def buscar_por_inciso_exacto(texto_completo, letra_inciso, expresion_esperada_norm, conjunto_esperado):
    lineas = texto_completo.split('\n')
    
    patrones_inciso = [
        rf"^{letra_inciso}[\)\.]",
        rf"\b{letra_inciso}[\)\.]",
        rf"inciso\s*{letra_inciso}\b"
    ]
    
    for i, linea in enumerate(lineas):
        linea_limpia = linea.strip()
        if not linea_limpia:
            continue
            
        contiene_inciso = any(re.search(patron, linea_limpia, re.IGNORECASE) 
                             for patron in patrones_inciso)
        
        if not contiene_inciso and i > 0:
            linea_anterior = lineas[i-1].strip()
            contiene_inciso = any(re.search(patron, linea_anterior, re.IGNORECASE) 
                                 for patron in patrones_inciso)
        
        if contiene_inciso:
            for j in range(i, min(i + 3, len(lineas))):
                linea_a_evaluar = lineas[j].strip()
                
                match_ecuacion = re.search(r"([^=]+)=\s*([^=]+)", linea_a_evaluar)
                if match_ecuacion:
                    expresion_encontrada = normalizar_expresion(match_ecuacion.group(1).strip())
                    conjunto_encontrado = extraer_conjunto(match_ecuacion.group(2).strip())
                    
                    if (expresion_encontrada == expresion_esperada_norm and 
                        conjunto_encontrado == conjunto_esperado):
                        return True, linea_a_evaluar
    
    return False, ""

def buscar_por_expresion_flexible(texto_completo, expresion_esperada_norm, conjunto_esperado):
    lineas = texto_completo.split('\n')
    
    for linea in lineas:
        linea_limpia = linea.strip()
        if not linea_limpia or '=' not in linea_limpia:
            continue
            
        match_ecuacion = re.search(r"([^=]+)=\s*([^=]+)", linea_limpia)
        if match_ecuacion:
            expresion_encontrada = normalizar_expresion(match_ecuacion.group(1).strip())
            conjunto_encontrado = extraer_conjunto(match_ecuacion.group(2).strip())
            
            if (expresion_encontrada == expresion_esperada_norm and 
                conjunto_encontrado == conjunto_esperado):
                return True, linea_limpia
    
    return False, ""

def determinar_videos_necesarios(indices_incorrectos):
    videos = []
    if 6 in indices_incorrectos:
        videos.append("https://youtu.be/-IHf20iF3Cg")
    
    otros_incorrectos = [i for i in indices_incorrectos if i != 6]
    if otros_incorrectos:
        videos.append("https://youtu.be/q5uYIWw7uD0")
    
    return videos

def mostrar_r3md():
    st.title("🔢 R3MD - Generador de retroalimentación por ejercicios de conjuntos")
    
    if not PDF_AVAILABLE:
        st.warning("⚠️ Las librerías de PDF no están instaladas. Solo se podrán procesar archivos Word (.docx)")
        st.info("Para habilitar soporte PDF, instala: pip install PyPDF2 pdfplumber")

    mensajes_exito = [
        "Excelente trabajo, {nombre}. El último ejercicio de este reto demuestra tu dominio de los conjuntos. Saludos.",
        "Muy bien hecho, {nombre}. Tus respuestas son precisas y completas. Sigue así.",
        "Perfecto, {nombre}. Se nota que comprendiste el tema de conjuntos.",
        "Buen trabajo, {nombre}. Has resuelto correctamente todos los incisos del reto.",
        "Todo correcto, {nombre}. Refleja que dominaste el concepto de operaciones con conjuntos.",
        "Felicidades, {nombre}. El ejercicio está resuelto sin errores.",
        "Gran resultado, {nombre}. El dominio del tema es evidente.",
        "Correcto en todos los puntos, {nombre}. Sigue con ese nivel.",
        "Buen cierre del reto, {nombre}. Todas las respuestas son válidas.",
        "Excelente resolución, {nombre}. Cada conjunto está trabajado con precisión."
    ]

    mensajes_error = [
        "Buen trabajo, {nombre}. Aunque hay detalles que revisar. Corrige y reenvía.",
        "Estás cerca, {nombre}. Revisa las operaciones que te señalo abajo y ajusta.",
        "Tu avance es bueno, {nombre}, pero hay expresiones que requieren corrección.",
        "Vamos bien, {nombre}, pero algunos incisos necesitan revisión.",
        "Buen intento, {nombre}, faltan ajustes en ciertas expresiones.",
        "Estás entendiendo el tema, {nombre}, pero hay errores por corregir.",
        "Revisa los conjuntos indicados abajo, {nombre}. Puedes mejorar.",
        "Vamos por buen camino, {nombre}, pero aún hay inconsistencias.",
        "Casi lo tienes, {nombre}. Corrige los puntos marcados como incorrectos.",
        "Un pequeño esfuerzo más, {nombre}, y todo estará correcto."
    ]

    EXPRESIONES_FIJAS = [
        "B ∩ C = {1,2,13}",
        "C ′ = {3,5,8,9,12,14}",
        "B ∪ C = {1,2,3,4,5,6,7,8,10,11,13}",
        "A ∩ C = {2,4,6,10}",
        "A ′ = {1,3,5,7,9,11,13}",
        "B – A = {1,3,5,13}",
        "C – B ′ = {1,2,13}"
    ]

    tipos_archivo = ["docx"]
    if PDF_AVAILABLE:
        tipos_archivo.append("pdf")

    documento_file = st.file_uploader(
        "Carga el archivo (Word .docx" + (" o PDF)" if PDF_AVAILABLE else " solamente)"), 
        type=tipos_archivo
    )

    col1, col2 = st.columns(2)
    with col1:
        usar_expresiones_fijas = st.checkbox("📋 Usar expresiones predefinidas", value=True)
    with col2:
        if not usar_expresiones_fijas:
            excel_file = st.file_uploader("📊 Carga archivo Excel personalizado", type=["xlsx"])
        else:
            excel_file = None

    if usar_expresiones_fijas:
        with st.expander("📝 Ver expresiones predefinidas que se evaluarán"):
            for i, expr in enumerate(EXPRESIONES_FIJAS):
                st.write(f"{chr(97+i)}) {expr}")
        cadenas_busqueda = EXPRESIONES_FIJAS
    else:
        cadenas_busqueda = []

    nombre = "Alumno"

    if documento_file:
        try:
            if documento_file.name.lower().endswith('.pdf'):
                if not PDF_AVAILABLE:
                    st.error("❌ No se pueden procesar archivos PDF. Instala las librerías necesarias: pip install PyPDF2 pdfplumber")
                    st.stop()
                st.info("📄 Procesando archivo PDF...")
                texto_completo = extraer_texto_pdf(documento_file)
            else:
                st.info("📄 Procesando archivo Word...")
                texto_completo = extraer_texto_docx(documento_file)

            normalizaciones = {
                'Ս': '∪',
                'Ո': '∩',
                '´': '′',
                '--': '–',
                ' - ': ' – ',
            }

            for original, reemplazo in normalizaciones.items():
                texto_completo = texto_completo.replace(original, reemplazo)

            nombre = extraer_nombre(texto_completo)
            
            with st.expander("👁️ Ver texto extraído (primeros 500 caracteres)"):
                st.text(texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo)

        except Exception as e:
            st.error(f"❌ Error leyendo el documento: {str(e)}")

    if documento_file and (usar_expresiones_fijas or excel_file):
        try:
            if not usar_expresiones_fijas and excel_file:
                df_cadenas = pd.read_excel(excel_file)
                columnas = df_cadenas.columns.tolist()
                columna_objetivo = st.selectbox("Selecciona la columna con las expresiones a buscar:", columnas)
                
                if columna_objetivo:
                    cadenas_busqueda = df_cadenas[columna_objetivo].astype(str).str.strip().unique().tolist()
                else:
                    cadenas_busqueda = []
            
            if cadenas_busqueda:
                coincidencias = []
                no_encontradas = []
                indices_incorrectos = []
                
                for i, expresion in enumerate(cadenas_busqueda):
                    encontrado, linea_encontrada = buscar_expresion_completa(texto_completo, i, expresion)
                    
                    if encontrado:
                        coincidencias.append(expresion)
                    else:
                        no_encontradas.append(expresion)
                        indices_incorrectos.append(i)

                st.success(f"✅ Total de expresiones a evaluar: {len(cadenas_busqueda)}")
                st.info(f"🎯 Coincidencias encontradas: {len(coincidencias)}")
                if no_encontradas:
                    st.warning(f"⚠️ No encontradas: {len(no_encontradas)}")

                mensaje_limpio = ""
                letras = "abcdefghijklmnopqrstuvwxyz"

                if len(no_encontradas) == 0:
                    encabezado = random.choice(mensajes_exito).format(nombre=nombre)
                    mensaje_limpio += f"{encabezado}\n\n"
                    for i, exp in enumerate(cadenas_busqueda):
                        mensaje_limpio += f"{letras[i]}) {exp} - correcto\n"
                else:
                    encabezado = random.choice(mensajes_error).format(nombre=nombre)
                    mensaje_limpio += f"{encabezado}\n"
                    
                    videos_necesarios = determinar_videos_necesarios(indices_incorrectos)
                    
                    if videos_necesarios:
                        mensaje_limpio += "Revisa estos videos:\n"
                        for video in videos_necesarios:
                            mensaje_limpio += f"{video}\n"
                        mensaje_limpio += "\n"
                    
                    for i, exp in enumerate(cadenas_busqueda):
                        if exp in coincidencias:
                            mensaje_limpio += f"{letras[i]}) {exp} - correcto\n"
                        else:
                            mensaje_limpio += f"{letras[i]}) - incorrecto\n"

                st.text_area("📝 Mensaje final generado para copiar:", value=mensaje_limpio, height=300)
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📋 Copiar al portapapeles", type="primary"):
                        components.html(copy_to_clipboard_js(mensaje_limpio), height=0)
                        st.success("✅ ¡Texto copiado al portapapeles!")
                
                with col2:
                    st.download_button("📥 Descargar mensaje como TXT", 
                                     data=mensaje_limpio, 
                                     file_name=f"retro_{nombre}.txt")

        except Exception as e:
            st.error(f"❌ Error al procesar los archivos: {str(e)}")

# ==================== R4MD - PROPOSICIONES LÓGICAS ====================

def buscar_columna_flexible(df, nombres_posibles):
    """
    Busca una columna de manera flexible, considerando diferentes variaciones de mayúsculas/minúsculas
    y espacios
    """
    columnas_df = df.columns.tolist()
    
    for nombre_buscado in nombres_posibles:
        # Búsqueda exacta
        if nombre_buscado in columnas_df:
            return nombre_buscado
        
        # Búsqueda insensible a mayúsculas/minúsculas
        for col in columnas_df:
            if col.lower() == nombre_buscado.lower():
                return col
        
        # Búsqueda con normalización de espacios
        nombre_normalizado = nombre_buscado.lower().strip()
        for col in columnas_df:
            col_normalizada = col.lower().strip()
            if col_normalizada == nombre_normalizado:
                return col
    
    return None

def mostrar_r4md():
    st.title("🧠 R4MD - Proposiciones Lógicas")
    
    mensajes_r4 = [
        "Buen día {nombre}. He tenido la oportunidad de revisar tu participación en el foro y quiero felicitarte, ya que has abordado todos los puntos de manera adecuada, cumpliendo con los criterios de la rúbrica. Ahora, aguardamos los comentarios de tus compañeros para enriquecer el intercambio. Te sugiero considerar sus observaciones y sacar provecho de esta oportunidad. ¡Saludos!",
        
        "Hola {nombre}, qué gusto saludarte. Revisé tu trabajo en el foro y quiero felicitarte por cumplir con los puntos solicitados en la rúbrica. Ahora esperemos la retroalimentación de tus compañeros, ya que el foro está diseñado para promover este intercambio de ideas. Aprovecha los comentarios recibidos para potenciar tu aprendizaje. Saludos.",
        
        "Gracias por tu aporte {nombre}. He revisado con detalle tu participación en el foro y quiero reconocerte el haber cumplido con todos los criterios establecidos. Ahora, esperamos las observaciones de tus compañeros, que enriquecerán la discusión y te brindarán nuevos puntos de vista. Aprovecha esta oportunidad para fortalecer tus conocimientos. Saludos cordiales.",
        
        "Excelente trabajo {nombre}. Al revisar tu contribución en el foro, pude ver que has cumplido con todos los aspectos solicitados en la rúbrica, ¡felicidades! Ahora queda por esperar los comentarios de tus compañeros, quienes podrán ofrecerte nuevas perspectivas. Considera sus observaciones para sacar el mayor provecho de esta actividad. Saludos.",
        
        "¿Qué tal? {nombre}. Muy bien hecho. Tu participación en el foro ha sido revisada, y es evidente que has cumplido con los puntos solicitados de forma satisfactoria. Ahora, espera la retroalimentación de tus compañeros, ya que el intercambio de ideas es el objetivo de este espacio. Aprovecha sus comentarios para fortalecer tu aprendizaje. ¡Saludos!"
    ]
    
    excel_file = st.file_uploader("📊 Carga el archivo Excel", type=["xlsx"])
    
    if excel_file:
        try:
            df = pd.read_excel(excel_file)
            
            # Mostrar información del archivo
            st.info(f"📋 Archivo cargado: {len(df)} filas, {len(df.columns)} columnas")
            
            # Mostrar columnas disponibles
            with st.expander("👁️ Ver columnas disponibles"):
                st.write(list(df.columns))
            
            # Buscar columnas de manera flexible
            nombres_columna_objetivo = [
                "Tarea:R4. Proposiciones lógicas (Real)",
                "Tarea: R4. Proposiciones lógicas (Real)",
                "Tarea:R4.Proposiciones lógicas (Real)"
            ]
            
            nombres_columna_nombre = [
                "Nombre",
                "nombre", 
                "NOMBRE",
                "Nombre completo",
                "nombre completo"
            ]
            
            columna_objetivo = buscar_columna_flexible(df, nombres_columna_objetivo)
            columna_nombre = buscar_columna_flexible(df, nombres_columna_nombre)
            
            if columna_objetivo:
                st.success(f"✅ Columna objetivo encontrada: '{columna_objetivo}'")
                
                # Filtrar filas con "-"
                filas_con_guion = df[df[columna_objetivo] == "-"]
                
                if len(filas_con_guion) > 0:
                    st.info(f"🔍 Encontradas {len(filas_con_guion)} filas con '-'")
                    
                    if columna_nombre:
                        st.success(f"✅ Columna nombre encontrada: '{columna_nombre}'")
                        
                        # Obtener nombres
                        nombres = filas_con_guion[columna_nombre].tolist()
                        
                        # Limpiar nombres (quitar espacios extra, NaN, etc.)
                        nombres_limpios = []
                        for nombre in nombres:
                            if pd.notna(nombre) and str(nombre).strip():
                                nombres_limpios.append(str(nombre).strip())
                        
                        if nombres_limpios:
                            # Crear mensajes balanceados
                            mensajes_finales = []
                            datos_para_excel = []
                            
                            st.markdown("---")
                            st.subheader("📝 Mensajes Generados")
                            
                            for i, nombre in enumerate(nombres_limpios):
                                # Usar módulo para distribuir mensajes de manera equilibrada
                                mensaje_idx = i % len(mensajes_r4)
                                mensaje_completo = mensajes_r4[mensaje_idx].format(nombre=nombre)
                                mensajes_finales.append(mensaje_completo)
                                
                                # Datos para Excel (nombre y mensaje en columnas separadas)
                                datos_para_excel.append({
                                    'Nombre': nombre,
                                    'Mensaje': mensaje_completo
                                })
                                
                                # Mostrar cada mensaje con su botón individual
                                with st.container():
                                    st.markdown(f"**{i+1}. {nombre}**")
                                    
                                    # Mostrar el mensaje en un área de texto pequeña
                                    st.text_area(
                                        f"Mensaje para {nombre}:", 
                                        value=mensaje_completo, 
                                        height=120, 
                                        key=f"mensaje_{i}",
                                        label_visibility="collapsed"
                                    )
                                    
                                    # Botón para copiar mensaje individual
                                    if st.button(f"📋 Copiar mensaje de {nombre}", key=f"copy_individual_{i}"):
                                        components.html(copy_to_clipboard_js(mensaje_completo), height=0)
                                        st.success(f"✅ ¡Mensaje de {nombre} copiado!")
                                    
                                    st.markdown("---")
                            
                            # Crear DataFrame para Excel con estructura solicitada
                            df_resultado = pd.DataFrame(datos_para_excel)
                            
                            st.success(f"✅ Procesados {len(mensajes_finales)} mensajes")
                            
                            # Mostrar DataFrame resultado
                            st.subheader("📊 Vista previa del Excel")
                            st.dataframe(df_resultado)
                            
                            # Botones principales
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                # Copiar todos los mensajes (solo el contenido, sin nombres)
                                texto_todos_mensajes = "\n\n".join(mensajes_finales)
                                if st.button("📋 Copiar TODOS los mensajes", type="primary"):
                                    components.html(copy_to_clipboard_js(texto_todos_mensajes), height=0)
                                    st.success("✅ ¡Todos los mensajes copiados!")
                            
                            with col2:
                                # Descargar Excel con estructura nombre|mensaje
                                output = io.BytesIO()
                                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                    df_resultado.to_excel(writer, index=False, sheet_name='Mensajes_R4')
                                
                                st.download_button(
                                    "📥 Descargar Excel",
                                    data=output.getvalue(),
                                    file_name="mensajes_r4.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                            
                            with col3:
                                # Descargar solo mensajes como TXT
                                st.download_button(
                                    "📄 Descargar mensajes TXT",
                                    data=texto_todos_mensajes,
                                    file_name="mensajes_r4.txt",
                                    mime="text/plain"
                                )
                            
                            # Mostrar distribución de mensajes
                            with st.expander("📊 Distribución de mensajes"):
                                distribucion = {}
                                for i in range(len(mensajes_finales)):
                                    mensaje_tipo = f"Mensaje {(i % len(mensajes_r4)) + 1}"
                                    distribucion[mensaje_tipo] = distribucion.get(mensaje_tipo, 0) + 1
                                
                                for tipo, cantidad in distribucion.items():
                                    st.write(f"{tipo}: {cantidad} veces")
                        
                        else:
                            st.warning("⚠️ No se encontraron nombres válidos en las filas con '-'")
                    
                    else:
                        st.error(f"❌ No se encontró ninguna columna de nombres")
                        st.write("**Columnas buscadas:** ", nombres_columna_nombre)
                        st.write("**Columnas disponibles:** ", list(df.columns))
                        
                        # Sugerir columnas similares
                        st.write("**💡 Sugerencias de columnas que podrían contener nombres:**")
                        for col in df.columns:
                            if any(palabra in col.lower() for palabra in ['nombre', 'name', 'alumno', 'estudiante']):
                                st.write(f"   - {col}")
                
                else:
                    st.warning("⚠️ No se encontraron filas con '-' en la columna objetivo")
                    
                    # Mostrar valores únicos de la columna objetivo para debug
                    with st.expander("🔍 Ver valores únicos en la columna objetivo"):
                        valores_unicos = df[columna_objetivo].value_counts()
                        st.write(valores_unicos)
            
            else:
                st.error(f"❌ No se encontró la columna objetivo")
                st.write("**Columnas buscadas:** ", nombres_columna_objetivo)
                st.write("**Columnas disponibles:** ", list(df.columns))
                
                # Sugerir columnas similares
                st.write("**💡 Sugerencias de columnas que podrían ser la objetivo:**")
                for col in df.columns:
                    if any(palabra in col.lower() for palabra in ['tarea', 'r4', 'proposiciones', 'logicas']):
                        st.write(f"   - {col}")
        
        except Exception as e:
            st.error(f"❌ Error al procesar el archivo Excel: {str(e)}")

# ==================== R7MD - MENSAJES PREDEFINIDOS ====================

def mostrar_r7md():
    st.title("💬 R7MD - Mensajes Predefinidos")
    
    # Mensajes para corregir
    mensajes_corregir = [
        """Buen trabajo, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 2, ya que en el paso 3, a pesar que identificas de manera correcta cada una de las relaciones transitivas, hay un cambio de dirección de la arista de "c" a "b", ya que la dirección en un paso anterior lo manejas de "b" a "c", de ahí la calificación, su pudieras argumentar dicho cambio de dirección podría corregir la calificación, quedo al pendiente.

Saludos.""",

        """Buen trabajo, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 2, ya que en el paso 3, no identificas en su totalidad las relaciones transitivas, situación que te lleva al error en tu diagrama final, te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio.

https://youtu.be/naYR2TQ84L0
Corrige y reenvía.

Saludos.""",

        """Buen trabajo, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 2, ya que en el paso 3, no entiendo como llegas a reestructurar para conseguir el diagrama presentado en el punto 4, ¿serías tan amable de dejarme una nota en tu reenvío? Te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio.

https://youtu.be/naYR2TQ84L0
Corrige y reenvía.

Saludos.""",

        """Buen trabajo, la primera tabla es correcta, en la parte que corresponde al dígrafo faltó eliminar la totalidad de las relaciones transitivas, hecho que no te permite alcanzar el 100% de la calificación.
Te dejo la resolución del ejercicio y quedo a disposición por si hubiera alguna duda más, aprovecho para preguntar, con todo respeto ¿Viste el video que te envíe en la realimentación anterior?
Saludos.

https://youtu.be/WTGkSBsLX34"""
    ]
    
    # Mensajes para correcto
    mensajes_correcto = [
        """Buen trabajo, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 3, ya que en el paso 4, estás realizando un acomodo incorrecto, situación que te lleva al error en tu diagrama final, te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio. Esperando tomes en consideración la recomendación, para evitar suspicacia en futuros trabajo, se asigna la mayor calificación.
https://youtu.be/WTGkSBsLX34

Éxito en tus subsecuentes retos.""",

        """Ha sido un placer acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Excelente trabajo, se requiere poner en práctica todo el conocimiento del curso para lograr resolver el ejercicio como lo has hecho, identificas de manera adecuada todos los elementos solicitados, continua así.

Ha sido un gusto acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos."""
    ]
    
    # Mostrar en dos columnas
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🔴 Mensajes para Corregir")
        for i, mensaje in enumerate(mensajes_corregir, 1):
            with st.expander(f"Mensaje {i} - Corregir"):
                st.text_area(f"Mensaje {i}", value=mensaje, height=200, key=f"corregir_{i}")
                if st.button(f"📋 Copiar Mensaje {i}", key=f"copy_corregir_{i}"):
                    components.html(copy_to_clipboard_js(mensaje), height=0)
                    st.success(f"✅ Mensaje {i} copiado!")
    
    with col2:
        st.subheader("🟢 Mensajes Correctos")
        for i, mensaje in enumerate(mensajes_correcto, 1):
            with st.expander(f"Mensaje {i} - Correcto"):
                st.text_area(f"Mensaje {i}", value=mensaje, height=200, key=f"correcto_{i}")
                if st.button(f"📋 Copiar Mensaje {i}", key=f"copy_correcto_{i}"):
                    components.html(copy_to_clipboard_js(mensaje), height=0)
                    st.success(f"✅ Mensaje {i} copiado!")
    
    # Botones para copiar todos los mensajes de cada categoría
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📋 Copiar TODOS los mensajes para Corregir", type="secondary"):
            todos_corregir = "\n\n" + "="*50 + "\n\n".join([f"MENSAJE {i+1} - CORREGIR:\n\n{msg}" for i, msg in enumerate(mensajes_corregir)])
            components.html(copy_to_clipboard_js(todos_corregir), height=0)
            st.success("✅ Todos los mensajes para corregir copiados!")
    
    with col2:
        if st.button("📋 Copiar TODOS los mensajes Correctos", type="secondary"):
            todos_correcto = "\n\n" + "="*50 + "\n\n".join([f"MENSAJE {i+1} - CORRECTO:\n\n{msg}" for i, msg in enumerate(mensajes_correcto)])
            components.html(copy_to_clipboard_js(todos_correcto), height=0)
            st.success("✅ Todos los mensajes correctos copiados!")

# ==================== NAVEGACIÓN PRINCIPAL ====================

if menu_option == "R3MD - Conjuntos":
    mostrar_r3md()
elif menu_option == "R4MD - Proposiciones Lógicas":
    mostrar_r4md()
elif menu_option == "R7MD - Mensajes Predefinidos":
    mostrar_r7md()
